import os
import configparser
import pandas as pd
import chromadb
import ollama
import langchain
from langchain.text_splitter import TokenTextSplitter
from docx import Document
from PyPDF2 import PdfReader

config = configparser.ConfigParser()
config.read('./literagbot.config')

class Corpus():
    def __init__(self):
        self.tables = {}
        self.texts = {}
        self.corpus_tables = []
        self.corpus_texts = []
        self.chunker = TokenTextSplitter(chunk_size=1000, chunk_overlap=100)
        self.db_ids = []
        self.db_metadatas = []
        self.db_docs = []
        self.chroma_client = chromadb.PersistentClient(path=config.get('CORPUS','STORE'))
        try:
            self.chroma_client.delete_collection(name=config.get('CORPUS','COLLECTION'))
        except:
            print(f"Collection {config.get('CORPUS','COLLECTION')} not found, nothing to clean!")
        self.collection = self.chroma_client.create_collection(name=config.get('CORPUS','COLLECTION'))


    def load_directory(self):
        for file in os.listdir(config.get('CORPUS','DATA_DIR')):
            filename = os.fsdecode(file)
            path = os.path.join(config.get('CORPUS','DATA_DIR'), filename)
            if filename.endswith(".csv"):
                print(f"Found CSV table: {filename}")
                table = pd.read_csv(path)
                self.tables[f'{filename}'] = table
            elif filename.endswith(".xlsx"):
                print(f"Found Excel table: {filename}")
                table = pd.read_excel(path)
                self.tables[f'{filename}'] = table
            elif filename.endswith(".txt"):
                print(f"Found text file: {filename}")
                text = open(path, "r", encoding="utf8")
                self.texts[f'{filename}'] = text.read()
            elif filename.endswith(".docx"):
                print(f"Found Word document: {filename}")
                document = Document(path)
                content = [p.text for p in document.paragraphs]
                self.texts[f'{filename}'] = ("\n".join(str(x) for x in content))
            elif filename.endswith(".pdf"):
                print(f"Found PDF: {filename}")
                reader = PdfReader(path)
                content = []
                for page in range(0, len(reader.pages)):
                    pageObject = reader.pages[page]
                    content.append(pageObject.extract_text)
                self.texts[f'{filename}'] = ("\n".join(str(x) for x in content))
        pass

    def chunk_tables(self):
        if config.get('CORPUS','SPLIT_2') is not None:
            for file in self.tables:
                table_chunks = {}
                table = self.tables[f'{file}']
                for category in table[config.get('CORPUS','SPLIT_1')]:
                    table_chunks[category] = {}
                for row in range(len(table)):
                    category = table[config.get('CORPUS','SPLIT_1')][row]
                    brand = table[config.get('CORPUS','SPLIT_2')][row]
                    if brand not in table_chunks[category]:
                        table_chunks[category][brand] = [table.loc[row]]
                    else:
                        table_chunks[category][brand].append([table.loc[row]])
                for category in table_chunks:
                    for brand in table_chunks[category]:
                        temp_docs = []
                        for doc in range(len(table_chunks[category][brand])):
                            for string in range(len(table_chunks[category][brand][doc])):
                                temp_docs.append(table_chunks[category][brand][0].index[string])
                                temp_docs.append(table_chunks[category][brand][doc][string])
                        self.corpus_tables.append(" ".join(str(x) for x in temp_docs))
        else:
            for file in self.tables:
                table_chunks = {}
                table = self.tables[f'{file}']
                for category in table[config.get('CORPUS','SPLIT_1')]:
                    table_chunks[category] = {}
                for category in table_chunks:
                    temp_docs = []
                    for doc in range(len(table_chunks[category])):
                        for string in range(len(table_chunks[category][brand][doc])):
                            temp_docs.append(table_chunks[category][0].index[string])
                            temp_docs.append(table_chunks[category][doc][string])
                    self.corpus_tables.append(" ".join(str(x) for x in temp_docs))
        pass

    def chunk_texts(self):
        for file in self.texts:
            split = self.chunker.split_text(self.texts[f'{file}'])
            for chunked_text in split:
                self.corpus_texts.append(chunked_text)
        pass

    def make_corpus(self):
        i = 0
        if len(self.corpus_tables) > 0:
            for doc in self.corpus_tables:
                self.db_ids.append(f"{i+1}")
                i+=1
                self.db_metadatas.append({'filetype':'table','file':'placeholder'})
                self.db_docs.append(doc)
        if len(self.corpus_texts) > 0:
            for doc in self.corpus_texts:
                self.db_ids.append(f"{i+1}")
                i+=1
                self.db_metadatas.append({'filetype':'text','file':'placeholder'})
                self.db_docs.append(doc)
        pass

    def load_vector_store(self):
        batch_size = 3000
        batch_count = (len(self.db_docs) - 1) // batch_size + 1

        for batch in range(0, batch_count):
            batch_docs = self.db_docs[batch*batch_size:(batch+1)*batch_size]
            batch_ids = self.db_ids[batch*batch_size:(batch+1)*batch_size]
            batch_metadatas = self.db_metadatas[batch*batch_size:(batch+1)*batch_size]
            print(f"Adding batch {batch+1}/{batch_count} to vector store")
            self.collection.add(documents=batch_docs, ids=batch_ids, metadatas=batch_metadatas)
            print(f"Finished adding batch {batch+1}/{batch_count} to vector store")
        pass


store = Corpus()
store.load_directory()
store.chunk_tables()
store.chunk_texts()
store.make_corpus()
store.load_vector_store()
ollama.pull(config.get('CHAT','MODEL'))
